from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "docs" / "nuxera-migration" / "docs" / "migration"
OUT_ROOT = SOURCE_DIR / "deliverables" / "2026-07-29"

SPANISH_DOCS = [
    "NUXERA_MANUAL_FUNCIONAL_DETALLADO_2026-07-29_ES.md",
    "NUXERA_MANUAL_TECNICO_DETALLADO_2026-07-29_ES.md",
    "NUXERA_PRUEBAS_Y_EVIDENCIA_VISUAL_2026-07-29_ES.md",
    "NUXERA_LO_QUE_FALTA_REALMENTE_2026-07-29_ES.md",
]

ENGLISH_DOCS = [
    "NUXERA_DETAILED_FUNCTIONAL_MANUAL_2026-07-29_EN.md",
    "NUXERA_DETAILED_TECHNICAL_MANUAL_2026-07-29_EN.md",
    "NUXERA_TESTING_AND_VISUAL_EVIDENCE_2026-07-29_EN.md",
    "NUXERA_REAL_REMAINING_WORK_2026-07-29_EN.md",
]


def set_run_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=9.5, bold=bold)


def add_summary_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(1.65)
    table.columns[1].width = Inches(4.85)
    set_cell_shading(table.rows[0].cells[0], "E8EEF5")
    set_cell_shading(table.rows[0].cells[1], "E8EEF5")
    set_cell_text(table.rows[0].cells[0], "Campo", bold=True)
    set_cell_text(table.rows[0].cells[1], "Detalle", bold=True)
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, bold=True)
        set_cell_text(cells[1], value)
    doc.add_paragraph()


def configure_doc(doc, title, subtitle, language_label):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.12

    for style_name, size, color in [
        ("Heading 1", 16, "1F4D78"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 11.5, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run(title)
    set_run_font(title_run, size=24, color="0B2545", bold=True)

    subtitle_p = doc.add_paragraph()
    subtitle_run = subtitle_p.add_run(subtitle)
    set_run_font(subtitle_run, size=11.5, color="555555")

    add_summary_table(
        doc,
        [
            ("Fecha", "2026-07-29"),
            ("Paquete", language_label),
            ("Producto", "NUXERA Financial Intelligence"),
            ("Uso", "Revision interna, socio, inversionistas y decision de cutover"),
        ],
    )


def add_footer(doc, label):
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(label)
    set_run_font(run, size=8.5, color="666666")


def add_image(doc, src, alt):
    image_path = (SOURCE_DIR / src).resolve()
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(6.3))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = caption.add_run(alt)
        set_run_font(cap_run, size=8.5, color="555555", italic=True)


def add_inline_markup(paragraph, text):
    parts = re.split(r"(`[^`]+`|\\*\\*[^*]+\\*\\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=9.5, color="0B2545")
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def build_docx(md_path, out_path, language_label):
    raw = md_path.read_text(encoding="utf-8").splitlines()
    title = next((line[2:].strip() for line in raw if line.startswith("# ")), md_path.stem)
    subtitle = "Documento generado desde la evidencia NUXERA" if language_label == "Espanol" else "Document generated from NUXERA evidence"

    doc = Document()
    configure_doc(doc, title, subtitle, language_label)
    add_footer(doc, f"NUXERA Financial Intelligence - {language_label} - 2026-07-29")

    pending_paragraph = []

    def flush_paragraph():
        nonlocal pending_paragraph
        if not pending_paragraph:
            return
        text = " ".join(pending_paragraph).strip()
        pending_paragraph = []
        if text:
            p = doc.add_paragraph()
            add_inline_markup(p, text)

    for line in raw:
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            continue
        if stripped.startswith("# "):
            continue
        if stripped.startswith("!["):
            flush_paragraph()
            match = re.match(r"!\\[(.*?)\\]\\((.*?)\\)", stripped)
            if match:
                add_image(doc, match.group(2), match.group(1))
            continue
        if stripped.startswith("### "):
            flush_paragraph()
            doc.add_paragraph(stripped[4:].strip(), style="Heading 3")
            continue
        if stripped.startswith("## "):
            flush_paragraph()
            doc.add_paragraph(stripped[3:].strip(), style="Heading 1")
            continue
        if stripped.startswith("- "):
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markup(p, stripped[2:].strip())
            continue
        if re.match(r"^\\d+\\.\\s+", stripped):
            flush_paragraph()
            p = doc.add_paragraph(style="List Number")
            add_inline_markup(p, re.sub(r"^\\d+\\.\\s+", "", stripped))
            continue
        pending_paragraph.append(stripped.rstrip("  "))

    flush_paragraph()
    doc.save(out_path)


def build_package(name, source_names, language_label):
    package_dir = OUT_ROOT / name
    if package_dir.exists():
        shutil.rmtree(package_dir)
    package_dir.mkdir(parents=True, exist_ok=True)
    for source_name in source_names:
        md_path = SOURCE_DIR / source_name
        docx_path = package_dir / source_name.replace(".md", ".docx")
        build_docx(md_path, docx_path, language_label)
    return package_dir


def main():
    spanish = build_package("NUXERA_Paquete_Espanol_2026-07-29", SPANISH_DOCS, "Espanol")
    english = build_package("NUXERA_English_Package_2026-07-29", ENGLISH_DOCS, "English")
    print(spanish)
    print(english)


if __name__ == "__main__":
    main()
